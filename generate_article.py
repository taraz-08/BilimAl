from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os

OUTPUT_DIR = r"c:\BilimAI"

# ─── DOCX ────────────────────────────────────────────────────────────────────

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_para(doc, text, bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def create_docx():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    t = doc.add_heading("BilimAI: Жасанды интеллект негізіндегі білім беру платформасы", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.font.color.rgb = RGBColor(37, 99, 235)
        run.font.size = Pt(16)

    doc.add_paragraph()

    # Authors
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Авторлар: Исламов Дінмұхамед, Ермаханбет Ерсін  |  Тобы: Delta  |  2026")
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()
    doc.add_horizontal_line = lambda: None  # placeholder

    # ABSTRACT
    add_heading(doc, "Аннотация", 1, (37, 99, 235))
    add_para(doc,
        "Қазіргі таңда жоғары білім беру жүйесінде студенттердің оқу үлгерімін бақылау, болжау және "
        "жеке тұлғаға бағытталған ұсынымдар беру мәселесі өте өзекті болып табылады. Осы мәселені шешу "
        "мақсатында BilimAI — жасанды интеллект технологияларына негізделген білім беру веб-платформасы "
        "әзірленді. Платформа студенттер мен оқытушыларға арналған кешенді цифрлық орта болып табылады.")

    # INTRO
    add_heading(doc, "Кіріспе", 1, (37, 99, 235))
    add_para(doc,
        "Дәстүрлі білім беру жүйесінде оқытушы бірнеше ондаған студенттің жеке академиялық жағдайын "
        "қадағалауға уақыт таба алмайды. Студент қиналып жатса да, бұл кеш байқалады. BilimAI платформасы "
        "осы олқылықты жою үшін жасалған — жасанды интеллект арқылы әрбір студентті жеке талдап, дер "
        "кезінде ескерту жасайды.")

    # STUDENT FEATURES
    add_heading(doc, "Студент үшін мүмкіндіктер", 1, (37, 99, 235))

    features = [
        ("1. Жеке кабинет және академиялық деректер",
         "Студент өзінің жеке кабинетіне кіріп, барлық академиялық ақпаратын бір жерден көре алады: "
         "қатысу пайызы, тест нәтижелері, бағалар, транскрипт."),
        ("2. Онлайн тесттер",
         "Студенттер платформа арқылы оқытушы жасаған тесттерді тапсырады. Тест нәтижелері автоматты "
         "түрде жүйеге сақталады."),
        ("3. Прокторинг жүйесі (Google Vision API)",
         "Тест тапсыру барысында Google Vision API технологиясы арқылы бейне бақылау жүзеге асырылады: "
         "бет анықтау, бас бұрылысын бақылау, телефон/кітап/ноутбук секілді тыйым салынған заттарды тану. "
         "Жүйе оқытушыға күдікті оқиғалар туралы есеп береді."),
        ("4. ИИ үлгерім болжамы (Gemini AI)",
         "Студенттің қатысу пайызы, тест ұпайлары, орташа баға талданып, академиялық жағдайы болжалады: "
         "\"Үздік\", \"Орташа\" немесе \"Қауіп аймағы\". Gemini AI қазақ тілінде жеке талдау жазады."),
        ("5. Жеке ұсынымдар (Gemini AI)",
         "Студенттің нақты жағдайына қарай Gemini AI қазақ тілінде 3 нақты кеңес береді: қай пәнге "
         "көбірек назар аудару керек, қатысуды қалай жақсарту керек және т.б."),
        ("6. Қолжазбаны тексеру",
         "Студент қолмен жазылған жауаптың фотосуретін жүктей алады, жүйе оны танып, мазмұнын талдайды."),
    ]
    for title, body in features:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(title + "\n")
        r.bold = True
        r.font.size = Pt(11)
        r2 = p.add_run(body)
        r2.font.size = Pt(11)

    # TEACHER FEATURES
    add_heading(doc, "Оқытушы үшін мүмкіндіктер", 1, (37, 99, 235))

    teacher_features = [
        ("1. Топтарды басқару", "Оқытушы топтар жасай алады, студенттерді топқа қоса алады."),
        ("2. Тест жасау", "Оқытушы кез келген пән бойынша тест жасайды: сұрақтар, нұсқалар, дұрыс жауаптар, уақыт шегі."),
        ("3. Студенттерді бақылау",
         "Оқытушы барлық студенттердің қатысу пайызын, тест нәтижелерін, ИИ болжамын және прокторинг "
         "есебін бір экраннан көре алады."),
        ("4. Бағалар мен транскрипт", "Оқытушы студенттерге баға қоя алады, транскрипт жазбаларын жүргізеді."),
    ]
    for title, body in teacher_features:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(title + "\n")
        r.bold = True
        r.font.size = Pt(11)
        r2 = p.add_run(body)
        r2.font.size = Pt(11)

    # TECH STACK TABLE
    add_heading(doc, "Техникалық шешімдер", 1, (37, 99, 235))
    table_data = [
        ["Компонент", "Технология"],
        ["Фронтенд", "React 18 + Vite + Tailwind CSS"],
        ["Бэкенд", "Python FastAPI"],
        ["Деректер қоры", "PostgreSQL (Supabase)"],
        ["ИИ болжам", "Google Gemini 2.5 Flash"],
        ["Прокторинг", "Google Vision API"],
        ["Аутентификация", "JWT токен"],
        ["Хостинг (фронтенд)", "Vercel"],
        ["Хостинг (бэкенд)", "Render.com"],
    ]
    table = doc.add_table(rows=len(table_data), cols=2)
    table.style = "Table Grid"
    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            run = cell.paragraphs[0].runs[0]
            if i == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                cell._tc.get_or_add_tcPr()
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "2563EB")
                shading.set(qn("w:color"), "auto")
                shading.set(qn("w:val"), "clear")
                cell._tc.tcPr.append(shading)
            run.font.size = Pt(10)

    # AI USAGE
    add_heading(doc, "Жасанды интеллект қалай қолданылады", 1, (37, 99, 235))
    ai_items = [
        ("Болжам — ", "студенттің деректері негізінде машиналық оқыту моделі үлгерім деңгейін анықтайды, ал Gemini оны қазақша түсіндіреді."),
        ("Ұсыным — ", "Gemini студенттің нақты олқылықтарын талдап, жеке кеңес береді."),
        ("Прокторинг — ", "Google Vision API суреттен адам бетін, бас бұрылысын және заттарды нақты анықтайды."),
    ]
    for bold_part, rest in ai_items:
        p = doc.add_paragraph(style="List Number")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(bold_part)
        r.bold = True
        r.font.size = Pt(11)
        r2 = p.add_run(rest)
        r2.font.size = Pt(11)

    # CONCLUSION
    add_heading(doc, "Қорытынды", 1, (37, 99, 235))
    add_para(doc,
        "BilimAI платформасы — қазақстандық жоғары оқу орындарына арналған, жасанды интеллект "
        "мүмкіндіктерін толыққанды пайдаланатын заманауи білім беру жүйесі. Платформа оқытушы мен "
        "студент арасындағы академиялық байланысты цифрландырады, үлгерім болжамы арқылы тәуекел "
        "аймағындағы студенттерді ерте анықтауға мүмкіндік береді, ал жеке ұсынымдар арқылы оқу "
        "сапасын арттырады.")
    add_para(doc,
        "Болашақта платформаны мобильді қосымша форматында шығару, жасанды интеллект мүмкіндіктерін "
        "кеңейту және Қазақстанның барлық жоғары оқу орындарына бейімдеу жоспарланып отыр.")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BilimAI — білімің болашағы бүгін басталады.")
    r.bold = True
    r.italic = True
    r.font.color.rgb = RGBColor(37, 99, 235)
    r.font.size = Pt(12)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Сайт: https://bilim-al-n6x7.vercel.app")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(100, 100, 100)

    path = os.path.join(OUTPUT_DIR, "BilimAI_Makala.docx")
    doc.save(path)
    print(f"DOCX saved: {path}")

# ─── PDF ─────────────────────────────────────────────────────────────────────

def create_pdf():
    path = os.path.join(OUTPUT_DIR, "BilimAI_Makala.pdf")
    doc = SimpleDocTemplate(
        path, pagesize=A4,
        rightMargin=2*cm, leftMargin=2.5*cm,
        topMargin=2*cm, bottomMargin=2*cm
    )

    styles = getSampleStyleSheet()

    BLUE = colors.HexColor("#2563EB")
    GRAY = colors.HexColor("#6B7280")
    LIGHT_BLUE = colors.HexColor("#EFF6FF")

    title_style = ParagraphStyle("title", fontSize=17, textColor=BLUE,
                                  alignment=TA_CENTER, spaceAfter=4, fontName="Helvetica-Bold", leading=22)
    subtitle_style = ParagraphStyle("subtitle", fontSize=10, textColor=GRAY,
                                     alignment=TA_CENTER, spaceAfter=12, fontName="Helvetica-Oblique")
    h1_style = ParagraphStyle("h1", fontSize=13, textColor=BLUE,
                               fontName="Helvetica-Bold", spaceBefore=14, spaceAfter=6, leading=18)
    body_style = ParagraphStyle("body", fontSize=10.5, alignment=TA_JUSTIFY,
                                 fontName="Helvetica", leading=16, spaceAfter=8)
    bold_body = ParagraphStyle("boldbody", fontSize=10.5, alignment=TA_JUSTIFY,
                                fontName="Helvetica-Bold", leading=16, spaceAfter=2)
    italic_center = ParagraphStyle("italic_c", fontSize=12, alignment=TA_CENTER,
                                    fontName="Helvetica-BoldOblique", textColor=BLUE, spaceAfter=4)
    footer_style = ParagraphStyle("footer", fontSize=9, alignment=TA_CENTER,
                                   textColor=GRAY, fontName="Helvetica-Oblique")

    story = []

    story.append(Paragraph("BilimAI: Жасанды интеллект негізіндегі білім беру платформасы", title_style))
    story.append(Paragraph("Авторлар: Исламов Дінмұхамед, Ермаханбет Ерсін  |  Тобы: Delta  |  2026", subtitle_style))
    story.append(HRFlowable(width="100%", thickness=1, color=BLUE, spaceAfter=10))

    def section(title, paragraphs):
        story.append(Paragraph(title, h1_style))
        for p in paragraphs:
            story.append(Paragraph(p, body_style))

    section("Аннотация", [
        "Қазіргі таңда жоғары білім беру жүйесінде студенттердің оқу үлгерімін бақылау, болжау және "
        "жеке тұлғаға бағытталған ұсынымдар беру мәселесі өте өзекті болып табылады. Осы мәселені шешу "
        "мақсатында <b>BilimAI</b> — жасанды интеллект технологияларына негізделген білім беру веб-платформасы "
        "әзірленді. Платформа студенттер мен оқытушыларға арналған кешенді цифрлық орта болып табылады."
    ])

    section("Кіріспе", [
        "Дәстүрлі білім беру жүйесінде оқытушы бірнеше ондаған студенттің жеке академиялық жағдайын "
        "қадағалауға уақыт таба алмайды. Студент қиналып жатса да, бұл кеш байқалады. BilimAI платформасы "
        "осы олқылықты жою үшін жасалған — жасанды интеллект арқылы әрбір студентті жеке талдап, дер "
        "кезінде ескерту жасайды."
    ])

    story.append(Paragraph("Студент үшін мүмкіндіктер", h1_style))
    student_features = [
        ("1. Жеке кабинет және академиялық деректер",
         "Студент өзінің жеке кабинетіне кіріп, барлық академиялық ақпаратын бір жерден көре алады: қатысу пайызы, тест нәтижелері, бағалар, транскрипт."),
        ("2. Онлайн тесттер",
         "Студенттер платформа арқылы оқытушы жасаған тесттерді тапсырады. Тест нәтижелері автоматты түрде жүйеге сақталады."),
        ("3. Прокторинг жүйесі (Google Vision API)",
         "Тест тапсыру барысында Google Vision API арқылы бейне бақылау жүзеге асырылады: бет анықтау, бас бұрылысын бақылау, телефон/кітап/ноутбук секілді тыйым салынған заттарды тану. Жүйе оқытушыға күдікті оқиғалар туралы есеп береді."),
        ("4. ИИ үлгерім болжамы (Gemini AI)",
         "Студенттің қатысу пайызы, тест ұпайлары, орташа баға талданып, академиялық жағдайы болжалады: \"Үздік\", \"Орташа\" немесе \"Қауіп аймағы\". Gemini AI қазақ тілінде жеке талдау жазады."),
        ("5. Жеке ұсынымдар (Gemini AI)",
         "Студенттің нақты жағдайына қарай Gemini AI қазақ тілінде 3 нақты кеңес береді: қай пәнге көбірек назар аудару керек, қатысуды қалай жақсарту керек және т.б."),
        ("6. Қолжазбаны тексеру",
         "Студент қолмен жазылған жауаптың фотосуретін жүктей алады, жүйе оны танып, мазмұнын талдайды."),
    ]
    for title, body in student_features:
        story.append(Paragraph(f"<b>{title}</b>", body_style))
        story.append(Paragraph(body, body_style))

    story.append(Paragraph("Оқытушы үшін мүмкіндіктер", h1_style))
    teacher_features = [
        ("1. Топтарды басқару", "Оқытушы топтар жасай алады, студенттерді топқа қоса алады."),
        ("2. Тест жасау", "Оқытушы кез келген пән бойынша тест жасайды: сұрақтар, нұсқалар, дұрыс жауаптар, уақыт шегі."),
        ("3. Студенттерді бақылау",
         "Оқытушы барлық студенттердің қатысу пайызын, тест нәтижелерін, ИИ болжамын және прокторинг есебін бір экраннан көре алады."),
        ("4. Бағалар мен транскрипт", "Оқытушы студенттерге баға қоя алады, транскрипт жазбаларын жүргізеді."),
    ]
    for title, body in teacher_features:
        story.append(Paragraph(f"<b>{title}</b>", body_style))
        story.append(Paragraph(body, body_style))

    story.append(Paragraph("Техникалық шешімдер", h1_style))
    table_data = [
        ["Компонент", "Технология"],
        ["Фронтенд", "React 18 + Vite + Tailwind CSS"],
        ["Бэкенд", "Python FastAPI"],
        ["Деректер қоры", "PostgreSQL (Supabase)"],
        ["ИИ болжам", "Google Gemini 2.5 Flash"],
        ["Прокторинг", "Google Vision API"],
        ["Аутентификация", "JWT токен"],
        ["Хостинг (фронтенд)", "Vercel"],
        ["Хостинг (бэкенд)", "Render.com"],
    ]
    t = Table(table_data, colWidths=[7*cm, 10*cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), BLUE),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("BACKGROUND", (0, 1), (-1, -1), LIGHT_BLUE),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, LIGHT_BLUE]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E1")),
        ("PADDING", (0, 0), (-1, -1), 7),
        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(t)

    section("Жасанды интеллект қалай қолданылады", [
        "<b>1. Болжам</b> — студенттің деректері негізінде машиналық оқыту моделі үлгерім деңгейін анықтайды, ал Gemini AI оны қазақша түсіндіреді.",
        "<b>2. Ұсыным</b> — Gemini AI студенттің нақты олқылықтарын талдап, жеке кеңес береді.",
        "<b>3. Прокторинг</b> — Google Vision API суреттен адам бетін, бас бұрылысын және заттарды нақты анықтайды.",
    ])

    section("Қорытынды", [
        "BilimAI платформасы — қазақстандық жоғары оқу орындарына арналған, жасанды интеллект "
        "мүмкіндіктерін толыққанды пайдаланатын заманауи білім беру жүйесі. Платформа оқытушы мен "
        "студент арасындағы академиялық байланысты цифрландырады, үлгерім болжамы арқылы тәуекел "
        "аймағындағы студенттерді ерте анықтауға мүмкіндік береді, ал жеке ұсынымдар арқылы оқу "
        "сапасын арттырады.",
        "Болашақта платформаны мобильді қосымша форматында шығару, жасанды интеллект мүмкіндіктерін "
        "кеңейту және Қазақстанның барлық жоғары оқу орындарына бейімдеу жоспарланып отыр.",
    ])

    story.append(Spacer(1, 0.3*cm))
    story.append(HRFlowable(width="100%", thickness=1, color=BLUE, spaceAfter=8))
    story.append(Paragraph("BilimAI — білімің болашағы бүгін басталады.", italic_center))
    story.append(Paragraph("https://bilim-al-n6x7.vercel.app", footer_style))

    doc.build(story)
    print(f"PDF saved: {path}")


create_docx()
create_pdf()
print("Done!")
